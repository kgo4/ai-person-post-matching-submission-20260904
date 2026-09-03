from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

people = [
('张晨','男 | 29岁 | 北京','Java后端 / AI平台工程师','5年','负责企业级AI平台与微服务架构，擅长Java、Spring Boot、向量检索和模型服务工程化。',
 ['Java、Spring Boot、Spring Cloud、MySQL、Redis、Kafka','Python、FastAPI、LangChain、Milvus、RAG','Docker、Kubernetes、Prometheus、GitLab CI'],
 ['AI知识库平台：设计文档解析、向量检索与权限体系，支撑日均30万次检索','模型服务网关：使用Java实现多模型路由、限流、重试和审计，降低接口故障率35%','负责团队代码评审、技术方案设计和线上故障排查']),
('李雨桐','女 | 27岁 | 上海','Java智能系统开发工程师','4年','聚焦Java服务开发与机器学习应用落地，具备从数据处理到模型调用的完整项目经验。',
 ['Java、Spring Boot、MyBatis-Plus、PostgreSQL','Python、PyTorch、Transformers、ONNX Runtime','Elasticsearch、RabbitMQ、Docker、Linux'],
 ['智能客服系统：构建Java业务服务和意图识别接口，日均处理12万次会话','推荐服务：完成特征查询、模型推理和结果缓存，接口P95延迟低于180ms','参与模型量化部署、灰度发布与监控告警建设']),
('王志远','男 | 32岁 | 深圳','高级Java / 大模型应用架构师','8年','负责高并发Java系统和大模型应用架构，擅长Agent工作流、知识图谱及平台治理。',
 ['Java、Spring Cloud Alibaba、Netty、DDD','大模型API、Agent工作流、RAG、Neo4j','MySQL、Redis、RocketMQ、Kubernetes'],
 ['企业Agent平台：设计多智能体编排、上下文隔离和结构化JSON输出校验','岗位能力图谱：构建岗位、能力、证据关系模型，支持动态演化分析','主导服务拆分、容量评估和线上稳定性治理，保障99.95%可用性']),
('陈思涵','女 | 25岁 | 杭州','Java AI应用开发工程师','3年','专注Java后端、智能问答和数据工程，能够快速将AI能力集成到业务流程。',
 ['Java、Spring Boot、JPA、Vue协作开发','Python、RAG、Embedding、OCR、Prompt工程','SQL、Redis、RabbitMQ、Docker'],
 ['合同审查助手：集成OCR与RAG检索，支持条款定位和风险提示','数据处理平台：开发异步任务、失败重试和批量导入能力','编写接口测试与性能测试，推动服务规范化交付']),
('赵启明','男 | 30岁 | 广州','Java智能制造平台工程师','6年','面向物联网与智能制造场景开发Java平台，熟悉实时数据处理和AI预测应用。',
 ['Java、Spring Boot、Netty、MQTT、WebSocket','Flink、Kafka、Python、机器学习模型服务','MySQL、InfluxDB、Redis、Kubernetes'],
 ['设备数据平台：接入2万台设备实时数据，构建告警与状态追踪服务','预测性维护：整合时序特征和模型推理接口，提前发现设备异常','负责平台性能优化、数据一致性设计和生产环境运维']),
]

def setup(doc):
 sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Pt(50); sec.left_margin=sec.right_margin=Pt(58)
 styles=doc.styles; styles['Normal'].font.name='Microsoft YaHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); styles['Normal'].font.size=Pt(10)
 return doc
doc=setup(Document()); sec=doc.sections[0]
styles=doc.styles; styles['Normal'].font.name='Microsoft YaHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); styles['Normal'].font.size=Pt(10)
title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=title.add_run('Java 与人工智能方向人员简历'); r.bold=True; r.font.size=Pt(20)
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.add_run('用于人员能力评估流程演示 | 生成日期：2026年8月31日').italic=True
for i,p in enumerate(people):
    if i: doc.add_section(WD_SECTION.NEW_PAGE)
    name,base,target,years,summary,skills,projects=p
    h=doc.add_paragraph(); rr=h.add_run(name); rr.bold=True; rr.font.size=Pt(18); h.add_run('  '+target).bold=True
    doc.add_paragraph(base+' | 工作经验：'+years)
    for label, text in [('职业概述',summary),('核心技能','；'.join(skills)),('项目经历','；'.join(projects))]:
        q=doc.add_paragraph(); q.paragraph_format.space_before=Pt(8); x=q.add_run(label); x.bold=True; x.font.size=Pt(12); doc.add_paragraph(text)
    doc.add_paragraph('教育背景：计算机科学与技术相关专业，本科').italic=True
doc.save(r'C:\Users\pengchao\Downloads\Java_AI_人员简历-汇总.docx')

for idx, group in enumerate(([people[0], people[1]], [people[2], people[3]], [people[4]],), 1):
    out = setup(Document())
    t = out.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER; z=t.add_run(f'Java 与人工智能方向人员简历（第{idx}份）'); z.bold=True; z.font.size=Pt(18)
    for j,p in enumerate(group):
        name,base,target,years,summary,skills,projects=p
        h=out.add_paragraph(); rr=h.add_run(name); rr.bold=True; rr.font.size=Pt(18); h.add_run('  '+target).bold=True
        out.add_paragraph(base+' | 工作经验：'+years)
        for label,text in [('职业概述',summary),('核心技能','；'.join(skills)),('项目经历','；'.join(projects))]:
            q=out.add_paragraph(); q.paragraph_format.space_before=Pt(8); x=q.add_run(label); x.bold=True; x.font.size=Pt(12); out.add_paragraph(text)
        out.add_paragraph('教育背景：计算机科学与技术相关专业，本科').italic=True
        if j < len(group)-1: out.add_page_break()
    out.save(rf'C:\Users\pengchao\Downloads\Java_AI_人员简历-{idx}.docx')
